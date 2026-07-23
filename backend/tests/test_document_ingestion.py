"""
test_document_ingestion.py — Loading/structure-extraction for the newer document
types (DOCX, HTML, CSV, XLSX). Uses real parser libraries against small fixture
files written to tmp_path — no mocking needed, these are fast, network-free parsers.
"""
import os

import pytest

import vectorstore as vectorstore_module
from vectorstore import LangChainRAG


@pytest.fixture(scope="module")
def rag(tmp_path_factory):
    # Module-scoped: LangChainRAG.__init__ loads the embedding model from disk, which is
    # slow to repeat per-test. None of these tests read/write via self.documents_folder —
    # they call the loader methods directly with an explicit filepath — so one shared
    # instance is safe. CHROMA_PERSIST_DIR ("./chroma_db" by default, i.e. the real
    # backend/chroma_db when tests run from backend/) is redirected to an isolated temp
    # dir so this never touches — or slowly reloads — the real production index, even
    # though none of these tests call create_vectorstore().
    from _pytest.monkeypatch import MonkeyPatch
    mp = MonkeyPatch()
    tmp_dir = tmp_path_factory.mktemp("docstore")
    mp.setattr(vectorstore_module, "CHROMA_PERSIST_DIR", str(tmp_dir / "_chroma"))
    instance = LangChainRAG(documents_folder=str(tmp_dir / "documents"))
    yield instance
    mp.undo()


def _write(tmp_path, name, content):
    path = tmp_path / name
    path.write_text(content, encoding="utf-8")
    return str(path)


class TestDocx:
    def test_load_docx_extracts_paragraph_text(self, rag, tmp_path):
        from docx import Document as DocxDocument

        docx_doc = DocxDocument()
        docx_doc.add_paragraph("Introduction paragraph about warranty terms.")
        docx_doc.add_paragraph("Second paragraph with more detail.")
        filepath = str(tmp_path / "policy.docx")
        docx_doc.save(filepath)

        docs = rag._load_docx(filepath, "policy.docx")
        assert len(docs) == 1
        assert "warranty terms" in docs[0].page_content
        assert "more detail" in docs[0].page_content
        assert docs[0].metadata == {"source": "policy.docx", "type": "docx"}

    def test_load_docx_skips_empty_file(self, rag, tmp_path):
        from docx import Document as DocxDocument

        filepath = str(tmp_path / "empty.docx")
        DocxDocument().save(filepath)

        assert rag._load_docx(filepath, "empty.docx") == []

    def test_extract_docx_structure_finds_headings(self, rag, tmp_path):
        from docx import Document as DocxDocument

        docx_doc = DocxDocument()
        docx_doc.add_paragraph("Chapter One", style="Heading 1")
        docx_doc.add_paragraph("Some body text.")
        docx_doc.add_paragraph("Chapter Two", style="Heading 1")
        filepath = str(tmp_path / "book.docx")
        docx_doc.save(filepath)

        struct = rag._extract_docx_structure(filepath, "book.docx")
        assert struct is not None
        assert "Chapter One" in struct.page_content
        assert "Chapter Two" in struct.page_content
        assert struct.metadata["chunk_type"] == "structure"

    def test_extract_docx_structure_none_below_two_headings(self, rag, tmp_path):
        from docx import Document as DocxDocument

        docx_doc = DocxDocument()
        docx_doc.add_paragraph("Only Heading", style="Heading 1")
        docx_doc.add_paragraph("Body text.")
        filepath = str(tmp_path / "short.docx")
        docx_doc.save(filepath)

        assert rag._extract_docx_structure(filepath, "short.docx") is None


class TestHtml:
    def test_load_html_strips_markup_and_scripts(self, rag, tmp_path):
        html = """
        <html><head><style>body{color:red}</style></head>
        <body>
            <script>console.log('nope')</script>
            <h1>Warranty Policy</h1>
            <p>Coverage lasts for one year from purchase.</p>
        </body></html>
        """
        filepath = _write(tmp_path, "page.html", html)

        docs = rag._load_html(filepath, "page.html")
        assert len(docs) == 1
        text = docs[0].page_content
        assert "Coverage lasts for one year" in text
        assert "console.log" not in text
        assert "color:red" not in text
        assert docs[0].metadata == {"source": "page.html", "type": "html"}

    def test_load_html_skips_empty_body(self, rag, tmp_path):
        filepath = _write(tmp_path, "empty.html", "<html><body></body></html>")
        assert rag._load_html(filepath, "empty.html") == []

    def test_extract_html_structure_finds_headings(self, rag, tmp_path):
        html = "<html><body><h1>Intro</h1><p>text</p><h2>Details</h2></body></html>"
        filepath = _write(tmp_path, "doc.html", html)

        struct = rag._extract_html_structure(filepath, "doc.html")
        assert struct is not None
        assert "Intro" in struct.page_content
        assert "Details" in struct.page_content
        assert struct.metadata["chunk_type"] == "structure"


class TestTabular:
    def test_load_csv_flattens_rows_with_repeated_header(self, rag, tmp_path):
        csv_text = "name,role\nAlice,Engineer\nBob,Designer\n"
        filepath = _write(tmp_path, "team.csv", csv_text)

        docs = rag._load_tabular(filepath, "team.csv", ".csv")
        assert len(docs) == 1
        assert "name | role" in docs[0].page_content
        assert "Alice | Engineer" in docs[0].page_content
        assert "Bob | Designer" in docs[0].page_content
        assert docs[0].metadata == {"source": "team.csv", "type": "tabular"}

    def test_load_csv_windows_large_row_counts_into_multiple_chunks(self, rag, tmp_path):
        rows = ["name,value"] + [f"item{i},{i}" for i in range(45)]
        filepath = _write(tmp_path, "big.csv", "\n".join(rows))

        docs = rag._load_tabular(filepath, "big.csv", ".csv")
        # 45 data rows / 20 per chunk -> 3 chunks
        assert len(docs) == 3
        for doc in docs:
            assert doc.page_content.startswith("name | value")

    def test_load_csv_skips_header_only_file(self, rag, tmp_path):
        filepath = _write(tmp_path, "header_only.csv", "name,value\n")
        assert rag._load_tabular(filepath, "header_only.csv", ".csv") == []

    def test_load_xlsx_flattens_rows(self, rag, tmp_path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["name", "role"])
        ws.append(["Alice", "Engineer"])
        ws.append(["Bob", "Designer"])
        filepath = str(tmp_path / "team.xlsx")
        wb.save(filepath)

        docs = rag._load_tabular(filepath, "team.xlsx", ".xlsx")
        assert len(docs) == 1
        assert "name | role" in docs[0].page_content
        assert "Alice | Engineer" in docs[0].page_content
        assert docs[0].metadata == {"source": "team.xlsx", "type": "tabular"}


class TestChunkDispatch:
    def test_docx_and_html_use_prose_chunker(self, rag):
        from langchain_core.documents import Document
        long_doc = Document(
            page_content="\n\n".join([f"Paragraph number {i} with some content." for i in range(60)]),
            metadata={"source": "long.docx", "type": "docx"},
        )
        chunks = rag._chunk_documents([long_doc])
        assert len(chunks) > 1
        assert all(c.metadata["source"] == "long.docx" for c in chunks)

    def test_tabular_chunks_pass_through_unsplit(self, rag):
        from langchain_core.documents import Document
        tabular_doc = Document(page_content="a | b\n1 | 2", metadata={"source": "t.csv", "type": "tabular"})
        chunks = rag._chunk_documents([tabular_doc])
        assert chunks == [tabular_doc]

    def test_new_extensions_registered_in_supported_extensions(self):
        for ext in (".docx", ".html", ".csv", ".xlsx"):
            assert ext in LangChainRAG.SUPPORTED_EXTENSIONS
